from fastapi import APIRouter, UploadFile, File, HTTPException, Depends, Body
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from app.models import Note, User
from app.utils.security import get_current_user, get_db
from app.utils.rate_limit import check_ocr_rate_limit
from app.utils.usage_stats import record_ocr_success, record_ocr_failure
import httpx
import os
import uuid
import shutil
import time
import re
import html
from typing import Optional
import logging
import platform
from PIL import Image
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from pydantic import BaseModel
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT
from io import BytesIO
from pathlib import Path

try:
    import arabic_reshaper
    from bidi.algorithm import get_display
except ImportError:
    arabic_reshaper = None
    get_display = None

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/notes", tags=["notes"])

# Resize large images before OCR to save memory and gain ~3× capacity (e.g. 4000x3000 -> 1500px max)
OCR_MAX_WIDTH = int(os.getenv("OCR_MAX_WIDTH", "1500"))
# PaddleOCR can exceed 15s on cold start, Arabic, or large images; was causing misleading "Processing failed, retry."
OCR_REQUEST_TIMEOUT_SEC = float(os.getenv("OCR_REQUEST_TIMEOUT_SEC", "180"))


def resize_image_for_ocr(file_path: str) -> str:
    """Resize image to max OCR_MAX_WIDTH px width; overwrites file. Returns same path."""
    try:
        with Image.open(file_path) as img:
            w, h = img.size
            if w <= OCR_MAX_WIDTH:
                return file_path
            ratio = OCR_MAX_WIDTH / w
            new_size = (OCR_MAX_WIDTH, int(h * ratio))
            out = img.resize(new_size, Image.Resampling.LANCZOS)
            fmt = img.format or "JPEG"
            save_kw = {"optimize": True}
            if fmt.upper() in ("JPEG", "JPG"):
                save_kw["quality"] = 90
            out.save(file_path, **save_kw)
            logger.info(f"Resized image {w}x{h} -> {new_size[0]}x{new_size[1]}")
    except Exception as e:
        logger.warning(f"Resize skipped for {file_path}: {e}")
    return file_path

# OCR Service URL - use environment variable or default
OCR_SERVICE_URL = os.getenv("OCR_SERVICE_URL", "http://ocr-service:9000")
UPLOAD_FOLDER = os.getenv("UPLOAD_FOLDER", "uploads")
# Normalize upload folder path for Windows compatibility
UPLOAD_FOLDER = os.path.normpath(UPLOAD_FOLDER)
if platform.system() == 'Windows':
    UPLOAD_FOLDER = UPLOAD_FOLDER.replace('/', '\\')
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

ARABIC_REGEX = re.compile(r"[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]")


def _contains_arabic(text: str) -> bool:
    return bool(ARABIC_REGEX.search(text or ""))


def _resolve_arabic_font_paths() -> tuple[Optional[str], Optional[str]]:
    """
    Return (regular, bold) font paths with Arabic glyph support.
    Prioritizes bundled fonts under backend/assets/fonts, then common system fonts.

    Uses backend/ as the anchor directory (parents[2] from app/routers/notes.py).
    The Docker image copies only backend/ to /app, so parents[3] was wrong there
    (/backend/... inside the container) and no TTF was found → Helvetica → tofu glyphs for Arabic.
    """
    backend_root = Path(__file__).resolve().parents[2]
    candidates = [
        (
            backend_root / "assets" / "fonts" / "NotoNaskhArabic-Regular.ttf",
            backend_root / "assets" / "fonts" / "NotoNaskhArabic-Bold.ttf",
        ),
        (
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
            Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"),
        ),
        (
            Path("/mnt/c/Windows/Fonts/arial.ttf"),
            Path("/mnt/c/Windows/Fonts/arialbd.ttf"),
        ),
        (
            Path("C:/Windows/Fonts/arial.ttf"),
            Path("C:/Windows/Fonts/arialbd.ttf"),
        ),
    ]
    for regular_path, bold_path in candidates:
        if regular_path.exists():
            return str(regular_path), str(bold_path) if bold_path.exists() else None
    logger.warning(
        "No Arabic-capable TTF found (bundle backend/assets/fonts or install fonts-dejavu-core). "
        "PDF Arabic will show as missing glyphs."
    )
    return None, None


def _register_pdf_fonts() -> dict:
    """
    Register PDF fonts and return a font map.
    Falls back safely to Helvetica if no Arabic-capable TTF exists.
    """
    fonts = {
        "regular": "Helvetica",
        "bold": "Helvetica-Bold",
        "italic": "Helvetica-Oblique",
    }
    regular_path, bold_path = _resolve_arabic_font_paths()
    if not regular_path:
        return fonts
    try:
        if "SmartNotebookArabicRegular" not in pdfmetrics.getRegisteredFontNames():
            pdfmetrics.registerFont(TTFont("SmartNotebookArabicRegular", regular_path))
        fonts["regular"] = "SmartNotebookArabicRegular"
        fonts["italic"] = "SmartNotebookArabicRegular"
        if bold_path:
            if "SmartNotebookArabicBold" not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont("SmartNotebookArabicBold", bold_path))
            fonts["bold"] = "SmartNotebookArabicBold"
        else:
            fonts["bold"] = "SmartNotebookArabicRegular"
    except Exception as font_error:
        logger.warning(f"Arabic PDF font registration failed, falling back to Helvetica: {font_error}")
    return fonts


def _to_pdf_paragraph_text(raw_text: str) -> tuple[str, bool]:
    """
    Escape text for reportlab Paragraph and shape Arabic when possible.
    Returns (safe_text, is_arabic).
    """
    value = raw_text or ""
    is_arabic = _contains_arabic(value)
    if is_arabic and arabic_reshaper and get_display:
        try:
            value = get_display(arabic_reshaper.reshape(value))
        except Exception as rtl_error:
            logger.warning(f"Arabic shaping failed; using raw text: {rtl_error}")
    safe = html.escape(value)
    return safe, is_arabic


# Pydantic models for request validation
class WordExportRequest(BaseModel):
    text: str
    title: str = "Smart Notebook Export"


async def call_ocr_service(image_path: str) -> dict:
    """
    Call OCR service to extract text from image.
    
    Args:
        image_path: Path to the image file
        
    Returns:
        Dictionary with OCR results (text, lines, etc.)
    """
    try:
        with open(image_path, "rb") as image_file:
            files = {"file": image_file}
            timeout = httpx.Timeout(
                connect=20.0,
                read=OCR_REQUEST_TIMEOUT_SEC,
                write=120.0,
                pool=20.0,
            )
            async with httpx.AsyncClient(timeout=timeout) as client:
                response = await client.post(
                    f"{OCR_SERVICE_URL}/ocr",
                    files=files
                )
                response.raise_for_status()
                return response.json()
    except httpx.TimeoutException:
        logger.error(
            "OCR service timeout for image %s after %.0fs read limit",
            image_path,
            OCR_REQUEST_TIMEOUT_SEC,
        )
        raise HTTPException(
            status_code=504,
            detail=(
                f"OCR timed out after {int(OCR_REQUEST_TIMEOUT_SEC)}s. "
                "Try again (first run loads models), use a smaller photo, or increase OCR_REQUEST_TIMEOUT_SEC."
            ),
        )
    except httpx.HTTPStatusError as e:
        logger.error(f"OCR service error: {e.response.status_code} - {e.response.text}")
        raise HTTPException(
            status_code=502,
            detail=f"OCR service error: {str(e)}"
        )
    except Exception as e:
        logger.error(f"Error calling OCR service: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process image with OCR: {str(e)}"
        )


@router.post("/ocr", response_model=dict)
async def process_image_with_ocr(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Upload an image, extract text using OCR service, and save to database.
    
    Returns:
        {
            "id": note_id,
            "text": "extracted text",
            "lines": ["line1", "line2"],
            "image_path": "path/to/image",
            "user_id": user_id
        }
    """
    try:
        check_ocr_rate_limit(current_user.id)

        # Validate file type
        if not file.content_type or not file.content_type.startswith('image/'):
            raise HTTPException(
                status_code=400,
                detail="File must be an image (jpg, png, etc.)"
            )
        
        # Generate unique filename
        file_id = str(uuid.uuid4())
        file_ext = os.path.splitext(file.filename or 'image.jpg')[1] or '.jpg'
        filename = f"{file_id}{file_ext}"
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        # Normalize path for Windows compatibility
        file_path = os.path.normpath(file_path)
        if platform.system() == 'Windows':
            file_path = file_path.replace('/', '\\')
        
        # Save uploaded file
        try:
            with open(file_path, "wb") as buffer:
                shutil.copyfileobj(file.file, buffer)
            resize_image_for_ocr(file_path)

            # Call OCR service
            logger.info(f"Calling OCR service for image: {filename}")
            t0 = time.monotonic()
            ocr_result = await call_ocr_service(file_path)
            duration = time.monotonic() - t0
            record_ocr_success(current_user.id, duration)
            
            # Create note in database
            note = Note(
                user_id=current_user.id,
                text=ocr_result.get("full_text", ocr_result.get("text", "")),
                image_path=file_path,
                title=file.filename or "Untitled Note"
            )
            db.add(note)
            db.commit()
            db.refresh(note)
            
            return {
                "id": note.id,
                "text": note.text,
                "lines": ocr_result.get("lines", []),
                "full_text": ocr_result.get("full_text", note.text),
                "image_path": note.image_path,
                "user_id": note.user_id,
                "title": note.title,
                "created_at": note.created_at.isoformat() if note.created_at else None
            }
            
        except HTTPException:
            record_ocr_failure(current_user.id)
            raise
        except Exception as e:
            record_ocr_failure(current_user.id)
            # Clean up uploaded file on error
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except:
                    pass
            logger.error(f"Error processing image: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Failed to process image: {str(e)}"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Unexpected error in OCR endpoint: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error: {str(e)}"
        )


@router.get("/", response_model=list)
async def get_notes(
    skip: int = 0,
    limit: int = 100,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all notes for the current user."""
    notes = db.query(Note).filter(
        Note.user_id == current_user.id
    ).offset(skip).limit(limit).all()
    
    return [
        {
            "id": note.id,
            "title": note.title,
            "text": note.text,
            "image_path": note.image_path,
            "created_at": note.created_at.isoformat() if note.created_at else None,
            "updated_at": note.updated_at.isoformat() if note.updated_at else None
        }
        for note in notes
    ]


@router.get("/{note_id}", response_model=dict)
async def get_note(
    note_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific note by ID."""
    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()
    
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    return {
        "id": note.id,
        "title": note.title,
        "text": note.text,
        "image_path": note.image_path,
        "created_at": note.created_at.isoformat() if note.created_at else None,
        "updated_at": note.updated_at.isoformat() if note.updated_at else None
    }


@router.delete("/{note_id}", response_model=dict)
async def delete_note(
    note_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a note by ID."""
    note = db.query(Note).filter(
        Note.id == note_id,
        Note.user_id == current_user.id
    ).first()
    
    if not note:
        raise HTTPException(status_code=404, detail="Note not found")
    
    # Delete image file if it exists
    if note.image_path and os.path.exists(note.image_path):
        try:
            os.remove(note.image_path)
        except Exception as e:
            logger.warning(f"Failed to delete image file {note.image_path}: {str(e)}")
    
    db.delete(note)
    db.commit()
    
    return {"message": "Note deleted successfully", "id": note_id}


@router.post("/export/word")
async def export_to_word(
    request: WordExportRequest
):
    """
    Export OCR text to a Word document.
    
    Request body:
        {
            "text": "OCR text content",
            "title": "Document title (optional)"
        }
    
    Returns:
        Word document file for download
    """
    try:
        text = request.text
        title = request.title
        # Create a new Document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_paragraph()
        title_run = title_paragraph.add_run(title)
        title_run.font.size = Pt(18)
        title_run.bold = True
        title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add spacing
        doc.add_paragraph()
        
        # Add metadata
        date_paragraph = doc.add_paragraph()
        date_run = date_paragraph.add_run(f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        date_run.font.size = Pt(10)
        date_run.italic = True
        date_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        doc.add_paragraph()
        doc.add_paragraph("─" * 50)
        doc.add_paragraph()
        
        # Add OCR text content
        # Split by lines to preserve formatting
        lines = text.split('\n')
        for line in lines:
            if line.strip():  # Only add non-empty lines
                p = doc.add_paragraph(line)
                p.style = 'Normal'
                for run in p.runs:
                    run.font.size = Pt(12)
            else:
                doc.add_paragraph()  # Empty line for spacing
        
        # Generate unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"smart_notebook_{timestamp}.docx"
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        # Normalize path for Windows compatibility
        file_path = os.path.normpath(file_path)
        if platform.system() == 'Windows':
            file_path = file_path.replace('/', '\\')
        
        # Save document
        doc.save(file_path)
        
        logger.info(f"Word document created: {filename}")
        
        # Return file for download
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create Word document: {str(e)}"
        )


class NoteContent(BaseModel):
    title: str
    text: str

class BulkExportRequest(BaseModel):
    notes: list[NoteContent]  # Accept note content directly
    title: str = "Smart Notebook - Multiple Notes"
    format: str = "word"  # "word" or "pdf"


@router.post("/export/bulk")
async def export_bulk_notes(
    request: BulkExportRequest,
    db: Session = Depends(get_db)
):
    """
    Export multiple notes as a single Word or PDF document.
    
    Request body:
        {
            "notes": [
                {"title": "Note 1", "text": "Content..."},
                {"title": "Note 2", "text": "Content..."}
            ],
            "title": "My Notes Collection",
            "format": "word" or "pdf"
        }
    
    Returns:
        Combined document file for download
    """
    try:
        if not request.notes or len(request.notes) == 0:
            raise HTTPException(status_code=400, detail="No notes provided")
        
        # Combine all note text
        combined_text = ""
        for i, note in enumerate(request.notes, 1):
            combined_text += f"\n\n{'='*50}\n"
            combined_text += f"Note {i}: {note.title or 'Untitled'}\n"
            combined_text += f"{'='*50}\n\n"
            combined_text += note.text
            combined_text += "\n\n"
        
        if request.format == "pdf":
            # Export as PDF
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"smart_notebook_bulk_{timestamp}.pdf"
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            # Normalize path for Windows compatibility
            file_path = os.path.normpath(file_path)
            if platform.system() == 'Windows':
                file_path = file_path.replace('/', '\\')
            
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            fonts = _register_pdf_fonts()
            
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor='#000000',
                spaceAfter=20,
                alignment=TA_CENTER,
                fontName=fonts["bold"]
            )
            
            safe_title, _ = _to_pdf_paragraph_text(request.title)
            story.append(Paragraph(safe_title, title_style))
            story.append(Spacer(1, 20))
            
            body_style = ParagraphStyle(
                'BodyStyle',
                parent=styles['Normal'],
                fontSize=12,
                leading=16,
                textColor='#000000',
                fontName=fonts["regular"],
                alignment=TA_LEFT
            )

            body_rtl_style = ParagraphStyle(
                'BodyRTLStyle',
                parent=body_style,
                alignment=TA_RIGHT
            )
            
            lines = combined_text.split('\n')
            for line in lines:
                if line.strip():
                    safe_line, is_arabic_line = _to_pdf_paragraph_text(line)
                    story.append(Paragraph(safe_line, body_rtl_style if is_arabic_line else body_style))
                    story.append(Spacer(1, 6))
                else:
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            
            return FileResponse(
                path=file_path,
                filename=filename,
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
        else:
            # Export as Word (default)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            filename = f"smart_notebook_bulk_{timestamp}.docx"
            file_path = os.path.join(UPLOAD_FOLDER, filename)
            # Normalize path for Windows compatibility
            file_path = os.path.normpath(file_path)
            if platform.system() == 'Windows':
                file_path = file_path.replace('/', '\\')
            
            doc = Document()
            
            # Add title
            title_paragraph = doc.add_paragraph()
            title_run = title_paragraph.add_run(request.title)
            title_run.font.size = Pt(18)
            title_run.bold = True
            title_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            doc.add_paragraph()
            
            # Add combined content
            lines = combined_text.split('\n')
            for line in lines:
                if line.strip():
                    p = doc.add_paragraph(line)
                    for run in p.runs:
                        run.font.size = Pt(12)
                else:
                    doc.add_paragraph()
            
            doc.save(file_path)
            
            return FileResponse(
                path=file_path,
                filename=filename,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}"}
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error creating bulk export: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create bulk export: {str(e)}"
        )


@router.post("/export/pdf")
async def export_to_pdf(
    request: WordExportRequest  # Reuse same request model
):
    """
    Export OCR text to a PDF document.
    
    Request body:
        {
            "text": "OCR text content",
            "title": "Document title (optional)"
        }
    
    Returns:
        PDF document file for download
    """
    try:
        text = request.text
        title = request.title
        
        # Generate unique filename
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"smart_notebook_{timestamp}.pdf"
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        # Normalize path for Windows compatibility
        file_path = os.path.normpath(file_path)
        if platform.system() == 'Windows':
            file_path = file_path.replace('/', '\\')
        
        # Create PDF document
        doc = SimpleDocTemplate(file_path, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        fonts = _register_pdf_fonts()
        
        # Create custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='#000000',
            spaceAfter=20,
            alignment=TA_CENTER,
            fontName=fonts["bold"]
        )
        
        date_style = ParagraphStyle(
            'DateStyle',
            parent=styles['Normal'],
            fontSize=10,
            textColor='#666666',
            spaceAfter=20,
            alignment=TA_LEFT,
            fontName=fonts["italic"]
        )
        
        body_style = ParagraphStyle(
            'BodyStyle',
            parent=styles['Normal'],
            fontSize=12,
            leading=16,
            textColor='#000000',
            fontName=fonts["regular"],
            alignment=TA_LEFT
        )

        body_rtl_style = ParagraphStyle(
            'BodyRTLStyle',
            parent=body_style,
            alignment=TA_RIGHT
        )
        
        # Add title
        safe_title, _ = _to_pdf_paragraph_text(title)
        story.append(Paragraph(safe_title, title_style))
        story.append(Spacer(1, 12))
        
        # Add metadata
        date_text = f"Exported on: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        story.append(Paragraph(date_text, date_style))
        story.append(Spacer(1, 12))
        
        # Add separator line
        story.append(Paragraph("─" * 80, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Add OCR text content
        # Split by lines and preserve formatting
        lines = text.split('\n')
        for line in lines:
            if line.strip():
                safe_line, is_arabic_line = _to_pdf_paragraph_text(line)
                story.append(Paragraph(safe_line, body_rtl_style if is_arabic_line else body_style))
                story.append(Spacer(1, 6))
            else:
                story.append(Spacer(1, 12))  # Empty line for spacing
        
        # Build PDF
        doc.build(story)
        
        logger.info(f"PDF document created: {filename}")
        
        # Return file for download
        return FileResponse(
            path=file_path,
            filename=filename,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
        
    except Exception as e:
        logger.error(f"Error creating PDF document: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to create PDF document: {str(e)}"
        )
